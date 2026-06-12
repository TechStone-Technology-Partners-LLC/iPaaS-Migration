"""
Generate GLD_Compliance_Business_Review.docx from the markdown source.
Adds TechStone logo in header of every page.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR  = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO_PATH = os.path.join(BASE_DIR, "logos", "TechStone.png")
OUT_PATH  = os.path.join(BASE_DIR, "WebMethods", "Analysis", "GLD_Compliance_Business_Review.docx")

# ── Brand Colors ──────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1F, 0x4E, 0x79)   # dark navy – headings
BLUE       = RGBColor(0x2E, 0x74, 0xB5)   # mid blue – subheadings / table header
TEAL       = RGBColor(0x00, 0x70, 0xC0)   # accent
LIGHT_BLUE = RGBColor(0xDD, 0xEB, 0xF7)   # table header fill
ALT_ROW    = RGBColor(0xF2, 0xF7, 0xFD)   # alternate table row
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY  = RGBColor(0x40, 0x40, 0x40)
MID_GRAY   = RGBColor(0x80, 0x80, 0x80)
WARN_AMBER = RGBColor(0xFF, 0xC0, 0x00)
GREEN_BG   = RGBColor(0xE2, 0xEF, 0xDA)
AMBER_BG   = RGBColor(0xFF, 0xF2, 0xCC)
CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)


# ── XML helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Set cell background shading via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top='4', bottom='4', left='4', right='4',
                     color='1F4E79', inner_color='D0D8E4'):
    """Add cell borders."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, sz in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_table_borders(table, color='2E74B5'):
    """Apply uniform borders to every cell in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),   'single')
                el.set(qn('w:sz'),    '4')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), color)
                tcBorders.append(el)
            tcPr.append(tcBorders)


def add_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def hex_from_rgb(rgb: RGBColor) -> str:
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'


# ── Header with logo ──────────────────────────────────────────────────────────
def add_header(doc):
    section = doc.sections[0]
    section.header_distance = Cm(1.0)
    header = section.header
    header.is_linked_to_previous = False

    # clear default empty paragraph
    for p in header.paragraphs:
        p.clear()

    hdr_para = header.paragraphs[0]
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Logo run
    logo_run = hdr_para.add_run()
    logo_run.add_picture(LOGO_PATH, height=Inches(0.4))

    # Separator line under header
    pPr = hdr_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Footer with page number ───────────────────────────────────────────────────
def add_footer(doc):
    section     = doc.sections[0]
    footer      = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def make_run(text, size=8):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz');  sz.set(qn('w:val'), str(size * 2))
        color = OxmlElement('w:color'); color.set(qn('w:val'), '808080')
        rPr.append(sz); rPr.append(color)
        r.append(rPr)
        t = OxmlElement('w:t'); t.text = text
        r.append(t)
        return r

    # "Page " text
    footer_para._p.append(make_run("Page "))

    # PAGE field
    r_page = OxmlElement('w:r')
    fc_begin = OxmlElement('w:fldChar'); fc_begin.set(qn('w:fldCharType'), 'begin')
    r_page.append(fc_begin)
    footer_para._p.append(r_page)

    r_instr = OxmlElement('w:r')
    it = OxmlElement('w:instrText'); it.text = 'PAGE'; it.set(qn('xml:space'), 'preserve')
    r_instr.append(it)
    footer_para._p.append(r_instr)

    r_end = OxmlElement('w:r')
    fc_end = OxmlElement('w:fldChar'); fc_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    footer_para._p.append(r_end)

    footer_para._p.append(
        make_run("  |  GLD Compliance Adapter Services — Business Review  |  TechStone LLC  |  Confidential")
    )


# ── Document styles ───────────────────────────────────────────────────────────
def apply_doc_defaults(doc):
    style = doc.styles['Normal']
    style.font.name     = 'Calibri'
    style.font.size     = Pt(10.5)
    style.font.color.rgb = DARK_GRAY

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


def style_h1(para, text):
    para.clear()
    run = para.add_run(text)
    run.font.name      = 'Calibri'
    run.font.bold      = True
    run.font.size      = Pt(16)
    run.font.color.rgb = NAVY
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)
    # Bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)


def style_h2(para, text):
    para.clear()
    run = para.add_run(text)
    run.font.name      = 'Calibri'
    run.font.bold      = True
    run.font.size      = Pt(12)
    run.font.color.rgb = BLUE
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)


def style_body(para, text='', bold=False, italic=False, color=None):
    if text:
        run = para.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.bold   = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        else:
            run.font.color.rgb = DARK_GRAY
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.space_before = Pt(0)
    return para


def style_code(para, text):
    run = para.add_run(text)
    run.font.name    = 'Courier New'
    run.font.size    = Pt(8.5)
    run.font.color.rgb = DARK_GRAY
    # light grey background via paragraph shading
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F5F5F5')
    pPr.append(shd)
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.left_indent  = Inches(0.25)


def add_callout(doc, text: str, bg_hex: str, prefix: str = ''):
    """Coloured single-paragraph callout box."""
    para = doc.add_paragraph()
    if prefix:
        run = para.add_run(prefix + '  ')
        run.font.bold = True
        run.font.size = Pt(10.5)
    run2 = para.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = DARK_GRAY
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg_hex)
    pPr.append(shd)
    para.paragraph_format.left_indent  = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(6)
    return para


# ── Table helpers ─────────────────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None, header_bg='1F4E79', alt_row=True):
    """Build a formatted table. headers=list, rows=list of lists."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold      = True
        run.font.size      = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name      = 'Calibri'
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_bg(hdr_cells[i], header_bg)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        bg_hex = hex_from_rgb(ALT_ROW) if (alt_row and r_idx % 2 == 1) else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            # Bold the first cell if it's a key column
            is_bold = (c_idx == 0 and len(headers) <= 3)
            run = p.add_run(str(val))
            run.font.size  = Pt(10)
            run.font.name  = 'Calibri'
            run.font.bold  = is_bold
            run.font.color.rgb = DARK_GRAY
            set_cell_bg(cells[c_idx], bg_hex)

    set_table_borders(table)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacing after table
    return table


# ── Cover page ────────────────────────────────────────────────────────────────
def add_cover(doc):
    # Logo centred large on cover
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(60)
    r_logo = p_logo.add_run()
    r_logo.add_picture(LOGO_PATH, height=Inches(1.1))

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(28)
    rt = p_title.add_run("GLD Compliance Adapter Services")
    rt.font.name      = 'Calibri'
    rt.font.bold      = True
    rt.font.size      = Pt(24)
    rt.font.color.rgb = NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p_sub.add_run("Migration Business Review")
    rs.font.name      = 'Calibri'
    rs.font.bold      = False
    rs.font.size      = Pt(16)
    rs.font.color.rgb = BLUE

    # Horizontal rule
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rule.paragraph_format.space_before = Pt(14)
    pPr   = p_rule._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bot   = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '12')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2E74B5')
    pBdr.append(bot)
    pPr.append(pBdr)

    # Metadata block
    meta = [
        ("Integration",      "GLDComplianceAdapterServices → Boomi"),
        ("Original Platform","webMethods Integration Server 6.5 (KeyBank, 2008)"),
        ("Target Platform",  "Boomi AtomSphere"),
        ("Prepared by",      "TechStone LLC"),
        ("Review Status",    "Awaiting Approval"),
    ]
    for label, value in meta:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.space_before = Pt(8)
        r_lbl = pm.add_run(f"{label}:  ")
        r_lbl.font.name      = 'Calibri'
        r_lbl.font.bold      = True
        r_lbl.font.size      = Pt(11)
        r_lbl.font.color.rgb = NAVY
        r_val = pm.add_run(value)
        r_val.font.name      = 'Calibri'
        r_val.font.size      = Pt(11)
        r_val.font.color.rgb = DARK_GRAY

    # Confidential notice
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_conf.paragraph_format.space_before = Pt(40)
    rc = p_conf.add_run("CONFIDENTIAL — For Internal Review Only")
    rc.font.name      = 'Calibri'
    rc.font.size      = Pt(9)
    rc.font.italic    = True
    rc.font.color.rgb = MID_GRAY

    # Page break after cover
    doc.add_page_break()


# ── Section: What Does This Integration Do ────────────────────────────────────
def section_overview(doc):
    p = doc.add_paragraph(); style_h1(p, "1.  What Does This Integration Do?")

    add_callout(doc,
        "When a customer applies for a GLD product (loan, account, line of credit, etc.), "
        "the bank is legally required to run a compliance check through an external system "
        "called the CIU (Compliance Identity Unit). This integration is responsible for "
        "everything that happens in the Oracle database around that compliance check.",
        'DDEEFF')

    bullets = [
        "Saves the customer's information before the check is sent",
        "Archives the raw request for regulatory audit purposes",
        "Stores the reference number returned by the CIU system",
        "Records whether the check passed or failed",
        "Logs any errors that occur during the process",
        "Retrieves full customer context when a CIU response arrives",
        "Cleans up old records on a scheduled maintenance cycle",
    ]
    for b in bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        r_b = p_b.add_run(b)
        r_b.font.size      = Pt(10.5)
        r_b.font.name      = 'Calibri'
        r_b.font.color.rgb = DARK_GRAY

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(8)
    rn = p_note.add_run("Important: ")
    rn.font.bold      = True
    rn.font.color.rgb = NAVY
    rn.font.size      = Pt(10.5)
    rn2 = p_note.add_run(
        "This integration does not make the compliance decision. It records everything "
        "before, during, and after the decision made by the external CIU system."
    )
    rn2.font.size      = Pt(10.5)
    rn2.font.color.rgb = DARK_GRAY


# ── Section: Flow ─────────────────────────────────────────────────────────────
def section_flow(doc):
    p = doc.add_paragraph(); style_h1(p, "2.  Integration Flow — Step by Step")

    steps = [
        ("1", "Receive the Compliance Request",
         "The process starts when it receives a compliance check request containing "
         "the customer's personal and application details (name, address, SSN/TIN, DOB, product information)."),
        ("2", "Save Request to Database",
         "All 25 fields from the incoming request are written to the Oracle database. "
         "Oracle generates a unique Request ID used to track this compliance check throughout its lifecycle."),
        ("3", "Archive the Raw Request",
         "The full XML payload of the original request is saved to the database for regulatory audit. "
         "This is a write-only step — it produces no output used by the process."),
        ("4", "Send to External CIU System  ⚠",
         "The compliance check request is forwarded to the external CIU compliance system. "
         "CIU returns a reference number linking the external record to our internal record. "
         "Note: The CIU endpoint is not defined in this package — it must be provided separately."),
        ("5", "Save CIU Reference Number",
         "The CIU reference number returned in Step 4 is written back into the original "
         "request record so both records are permanently linked."),
        ("6", "Wait for the Compliance Result",
         "The CIU system processes the request and returns a pass or fail result. "
         "This response arrives asynchronously (the process pauses here until CIU responds)."),
        ("7", "Retrieve Full Customer Record",
         "Using the CIU reference number, the process looks up the complete customer "
         "and request details from the Oracle database."),
        ("8a", "Record the Result — PASS",
         "If the compliance check passed: the success result is logged with the Check Type "
         "and CIU Reference Number."),
        ("8b", "Record the Result — FAIL / ERROR",
         "If the compliance check failed or an error occurred: the error type, error code, "
         "and description are logged with the CIU Reference Number."),
        ("9", "Maintenance — Purge Old Records",
         "On a scheduled basis (independent of live checks), old compliance records are deleted "
         "from the Oracle database. The age threshold is controlled by the database stored procedure."),
    ]

    for num, title, desc in steps:
        # Step title
        p_step = doc.add_paragraph()
        p_step.paragraph_format.space_before = Pt(10)
        p_step.paragraph_format.space_after  = Pt(2)
        p_step.paragraph_format.left_indent  = Inches(0.25)
        r_num = p_step.add_run(f"Step {num}  ")
        r_num.font.bold      = True
        r_num.font.size      = Pt(11)
        r_num.font.color.rgb = BLUE
        r_num.font.name      = 'Calibri'
        r_title = p_step.add_run(title)
        r_title.font.bold      = True
        r_title.font.size      = Pt(11)
        r_title.font.color.rgb = NAVY
        r_title.font.name      = 'Calibri'

        # Step description
        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.left_indent  = Inches(0.5)
        p_desc.paragraph_format.space_after  = Pt(2)
        rd = p_desc.add_run(desc)
        rd.font.size      = Pt(10.5)
        rd.font.color.rgb = DARK_GRAY
        rd.font.name      = 'Calibri'


# ── Section: 7 Operations ─────────────────────────────────────────────────────
def section_operations(doc):
    p = doc.add_paragraph(); style_h1(p, "3.  The 7 Database Operations")

    p_intro = doc.add_paragraph()
    style_body(p_intro,
        "Each of the following operations is a call from Boomi into the Oracle database. "
        "The underlying database stored procedures remain unchanged — Boomi simply calls them "
        "with the same inputs webMethods used.")

    headers = ["#", "Operation", "What It Does", "Data In", "Data Out"]
    rows = [
        ("1", "Log Check Request",
         "Saves a new compliance check request to Oracle",
         "Customer name, address, DOB, SSN/TIN, product details (25 fields)",
         "Internal Request ID (auto-generated)"),
        ("2", "Log Check Request XML",
         "Saves the raw XML payload for audit trail",
         "Application ID, full XML string, 3 identifiers",
         "Nothing (write-only)"),
        ("3", "Log Check Reply",
         "Records the compliance check result",
         "CIU Reference Number, Check Type, Pass/Fail",
         "Nothing (write-only)"),
        ("4", "Log Check Reply Error",
         "Records error details if the check fails",
         "Error type, error code, description, CIU Reference",
         "Nothing (write-only)"),
        ("5", "Select Customer and Request",
         "Retrieves the full customer + request record",
         "CIU Reference Number",
         "All 28 customer + request fields"),
        ("6", "Update CIU Reference Number",
         "Links our internal record to the CIU external record",
         "Internal Request ID + CIU Reference Number",
         "Nothing (write-only)"),
        ("7", "Purge Data",
         "Deletes old/expired compliance records",
         "None (threshold built into Oracle procedure)",
         "Nothing (housekeeping)"),
    ]
    add_table(doc, headers, rows,
              col_widths=[0.3, 1.8, 2.0, 2.2, 1.8])


# ── Section: Data ─────────────────────────────────────────────────────────────
def section_data(doc):
    p = doc.add_paragraph(); style_h1(p, "4.  Data Handled by This Integration")

    p2 = doc.add_paragraph(); style_h2(p2, "4.1  Customer Information Sent to the Database")

    headers = ["Category", "Fields"]
    rows = [
        ("Identity",            "First Name, Middle Name, Last Name, Business Name"),
        ("Address",             "Address Lines 1–4, City, State (2 chars), ZIP, Country Code"),
        ("Sensitive / PII ⚠",  "SSN or Tax ID (SSNTIN),  Date of Birth (DOB)"),
        ("Contact",             "Compliance Reply Email"),
        ("Account",             "Customer Number, Customer Type, Party Type"),
        ("Product/Application", "Application Number, Channel, Line of Business, Product Code, Sub-Product Code, Post-Back URL"),
    ]
    add_table(doc, headers, rows, col_widths=[2.0, 5.0])

    p3 = doc.add_paragraph(); style_h2(p3, "4.2  Tracking Information Generated During the Process")

    headers2 = ["Field", "When It Appears", "What It Is"]
    rows2 = [
        ("Internal Request ID",    "After Step 2 — Save to Database",   "Oracle-generated unique ID for the compliance record"),
        ("CIU Reference Number",   "After Step 4 — CIU responds",        "Reference assigned by the external compliance system"),
        ("Compliance Result",      "After Step 6 — Result arrives",      "TRUE (passed) or FALSE (failed)"),
    ]
    add_table(doc, headers2, rows2, col_widths=[2.0, 2.5, 3.0])

    p4 = doc.add_paragraph(); style_h2(p4, "4.3  Database Tables")

    p_intro = doc.add_paragraph()
    style_body(p_intro, "Two Oracle tables are written to and read from by this integration:")

    p5 = doc.add_paragraph()
    p5.paragraph_format.space_before = Pt(6)
    r5 = p5.add_run("ACCCUSTOMER — Customer Master")
    r5.font.bold = True; r5.font.size = Pt(10.5); r5.font.color.rgb = NAVY

    headers3 = ["Column", "Meaning"]
    rows3 = [
        ("ACCCUSTOMERID",  "Unique customer ID (generated by Oracle)"),
        ("CUSTOMERNBR",    "Customer number from source system"),
        ("FIRSTNAME / MIDDLENAME / LASTNAME", "Customer full name"),
        ("ADDRESSLINE1–4, CITY, STATE, ZIP",  "Full mailing address"),
        ("SSNTIN",         "Social Security Number or Tax ID  — sensitive data"),
        ("DOB",            "Date of birth  — sensitive data"),
        ("CUSTOMERTYPE, PARTYTYPE, BUSINESSNAME", "Classification and business details"),
    ]
    add_table(doc, headers3, rows3, col_widths=[2.8, 4.7])

    p6 = doc.add_paragraph()
    r6 = p6.add_run("ACCCHECKREQUEST — Compliance Request Log")
    r6.font.bold = True; r6.font.size = Pt(10.5); r6.font.color.rgb = NAVY

    headers4 = ["Column", "Meaning"]
    rows4 = [
        ("ACCCHECKREQUESTID",  "Unique request ID (generated by Oracle)"),
        ("ACCCUSTOMERID",      "Links back to ACCCUSTOMER (foreign key)"),
        ("APPLICATIONNBR, CHANNEL, LOB", "What product/channel triggered the check"),
        ("PRODUCTCODE, SUBPRODUCTCODE",  "Product details"),
        ("CIUREFNBR",          "Filled in after CIU responds (empty on initial insert)"),
        ("REQUESTTIMESTAMP",   "Date and time when the check was submitted"),
    ]
    add_table(doc, headers4, rows4, col_widths=[2.8, 4.7])


# ── Section: Transformations ──────────────────────────────────────────────────
def section_transforms(doc):
    p = doc.add_paragraph(); style_h1(p, "5.  Data Transformations")

    p_intro = doc.add_paragraph()
    style_body(p_intro,
        "The large majority of fields pass through this integration unchanged (direct copy). "
        "Two fields require a format conversion:")

    headers = ["Field", "Change Made", "Reason"]
    rows = [
        ("Date of Birth (DOB)",
         "Text string 'yyyy-MM-dd'  →  Oracle DATE type",
         "Oracle's stored procedure expects a native Oracle date, not a text string. Boomi applies a built-in date conversion function."),
        ("Compliance Result",
         "Boolean true/false  →  Text 'TRUE' / 'FALSE'",
         "The Oracle procedure expects text; webMethods passed a Boolean. Boomi converts this before calling the database."),
    ]
    add_table(doc, headers, rows, col_widths=[1.8, 2.5, 3.5])


# ── Section: Risks ────────────────────────────────────────────────────────────
def section_risks(doc):
    p = doc.add_paragraph(); style_h1(p, "6.  Risks and Considerations")

    items = [
        ("⚠  Sensitive Data (PII)",
         "SSN/TIN and Date of Birth flow through the Boomi integration and are stored in Oracle. "
         "Boomi process logging will be disabled to prevent these values from appearing in integration audit logs. "
         "The database-level controls (encryption, access, masking) remain the same — Boomi calls the same Oracle stored procedures as webMethods did.",
         'FFF2CC'),
        ("⚠  External CIU System (Step 4)",
         "The external compliance system URL, credentials, and request format are not part of this package. "
         "In webMethods, this was handled by a separate companion package (GLDComplianceCheck). "
         "Without the CIU endpoint, Step 4 of the Boomi process will be built with a placeholder. "
         "This is the most significant open item.",
         'FFF2CC'),
        ("✅  Database Credentials",
         "The Oracle database password is not stored in the Boomi integration. "
         "It will be configured separately through Boomi's secure Environment Extensions — "
         "equivalent to a secrets vault.",
         'E2EFDA'),
        ("✅  Stored Procedure Logic",
         "The business logic inside each Oracle stored procedure (INSERT rules, validation, error handling) "
         "is not visible to the migration team — it runs entirely on the Oracle server. "
         "Boomi's role is to call the procedures with the correct inputs. The procedure bodies are unchanged.",
         'E2EFDA'),
    ]

    for title, body, bg in items:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(8)
        rt = p_title.add_run(title)
        rt.font.bold      = True
        rt.font.size      = Pt(11)
        rt.font.color.rgb = NAVY
        rt.font.name      = 'Calibri'
        add_callout(doc, body, bg)


# ── Section: Before vs After ──────────────────────────────────────────────────
def section_comparison(doc):
    p = doc.add_paragraph(); style_h1(p, "7.  Before vs. After Comparison")

    headers = ["Aspect", "webMethods (Before)", "Boomi (After)"]
    rows = [
        ("Platform",             "webMethods IS 6.5 (2008, decommissioning)",      "Boomi AtomSphere (current, cloud-native)"),
        ("Database Calls",       "7 JDBC Adapter Services",                         "7 DatabaseV2 Operations (same Oracle stored procedures)"),
        ("Connection",           "GLDComplianceAdapterEnv:ExpressOS alias",          "Boomi DatabaseV2 Connection component (same JDBC URL)"),
        ("PII Logging",          "Dependent on IS server logging config",            "Explicitly disabled in Boomi process config"),
        ("Credentials",          "Stored in webMethods adapter config",              "Stored in Boomi Environment Extensions (secure)"),
        ("Maintenance Purge",    "Scheduled IS trigger",                             "Boomi scheduled process (same SP called)"),
        ("Functionality",        "100% baseline",                                    "100% of database operations preserved"),
    ]
    add_table(doc, headers, rows,
              col_widths=[1.9, 2.5, 3.0],
              header_bg='2E74B5')


# ── Section: Sign-Off ─────────────────────────────────────────────────────────
def section_signoff(doc):
    p = doc.add_paragraph(); style_h1(p, "8.  Sign-Off Checklist")

    p_intro = doc.add_paragraph()
    style_body(p_intro,
        "The following items require a decision or confirmation before the Boomi build can be completed. "
        "Please review each item and indicate your response:")

    headers = ["#", "Item", "Owner", "Priority", "Status"]
    rows = [
        ("1", "CIU system endpoint URL, credentials, and request/response format — must be provided before Step 4 can be built",
         "Architecture / Business", "HIGH", "Open"),
        ("2", "Confirm target database host: is CSC06DSHORA1S:1522/ILMSUM correct for the migrated environment?",
         "Infrastructure / DBA", "High", "Open"),
        ("3", "Provide Oracle database password (via Boomi Environment Extensions — not written into the integration)",
         "DBA / Security", "High", "Open"),
        ("4", "Confirm that disabling Boomi audit logging for SSN/TIN fields meets compliance requirements",
         "Security / Compliance", "Medium", "Open"),
        ("5", "Confirm purge data date threshold: does ACCPURGEDATA retain the correct threshold for the new environment?",
         "DBA", "Low", "Open"),
    ]
    add_table(doc, headers, rows,
              col_widths=[0.3, 3.5, 1.6, 0.85, 0.85],
              header_bg='1F4E79')

    doc.add_paragraph()
    add_callout(doc,
        "Once all 5 items above are resolved, the 9 Boomi components can be built and tested. "
        "Post-build, the Oracle password configuration and CIU connector wiring will be completed "
        "in the Boomi UI by the developer.",
        'DDEEFF', prefix='Next Step:')


# ── Section: Known Gaps ───────────────────────────────────────────────────────
def section_gaps(doc):
    p = doc.add_paragraph(); style_h1(p, "9.  Known Gaps — Items Requiring Input")

    headers = ["#", "Gap", "Who Needs to Act", "Priority"]
    rows = [
        ("1", "CIU external service URL and credentials are unknown — blocks Step 4",
         "Business / Architecture", "HIGH"),
        ("2", "Stored procedure bodies not available — only the call signatures are known (SP logic runs on Oracle server unchanged)",
         "DBA team", "Medium"),
        ("3", "Database table DDL (exact schema definition) inferred from adapter metadata, not confirmed from DBA",
         "DBA team", "Medium"),
        ("4", "DOB input format — confirm it is always yyyy-MM-dd coming into the Boomi process",
         "Source system team", "Low"),
        ("5", "Purge data threshold — confirm what date range ACCPURGEDATA currently uses",
         "DBA team", "Low"),
    ]
    add_table(doc, headers, rows,
              col_widths=[0.3, 3.7, 1.8, 0.85])


# ── Build Document ─────────────────────────────────────────────────────────────
def build():
    doc = Document()
    apply_doc_defaults(doc)
    add_header(doc)
    add_footer(doc)

    add_cover(doc)
    section_overview(doc)
    section_flow(doc)
    doc.add_page_break()
    section_operations(doc)
    doc.add_page_break()
    section_data(doc)
    doc.add_page_break()
    section_transforms(doc)
    section_risks(doc)
    doc.add_page_break()
    section_comparison(doc)
    section_signoff(doc)
    section_gaps(doc)

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
