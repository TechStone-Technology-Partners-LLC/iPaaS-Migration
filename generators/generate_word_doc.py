"""
Migration Design Document Generator
Converts a migration spec JSON into a branded Word (.docx) design document.

Usage:
    python generators/generate_word_doc.py migration-specs/myproject.json
    python generators/generate_word_doc.py migration-specs/myproject.json --output docs/MyProject_Design.docx
    python generators/generate_word_doc.py migration-specs/myproject.json --md-dir WebMethods/MD/
"""

import argparse
import json
import os
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Brand Colors ──────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1F, 0x4E, 0x79)
BLUE       = RGBColor(0x2E, 0x74, 0xB5)
TEAL       = RGBColor(0x00, 0x70, 0xC0)
LIGHT_BLUE = RGBColor(0xDD, 0xEB, 0xF7)
ALT_ROW    = RGBColor(0xF2, 0xF7, 0xFD)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY  = RGBColor(0x40, 0x40, 0x40)
AMBER_BG   = RGBColor(0xFF, 0xF2, 0xCC)
GREEN_BG   = RGBColor(0xE2, 0xEF, 0xDA)
CODE_BG    = RGBColor(0xF5, 0xF5, 0xF5)

SEVERITY_COLORS = {
    "high":   "FFDDC1",
    "medium": "FFF2CC",
    "low":    "E2EFDA",
}


# ── XML helpers ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_table_borders(table, color="D0D8E4"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "bottom", "left", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), color)
                borders.append(el)
            tcPr.append(borders)


# ── Text helpers ──────────────────────────────────────────────────────────────
def _h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E74B5")
    pb.append(bot)
    pPr.append(pb)


def _h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


def _h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)


def _body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_after = Pt(3)
    return p


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)


def _callout(doc, text, bg_hex="FFF2CC", prefix=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    _set_cell_bg(cell, bg_hex)
    p = cell.paragraphs[0]
    run = p.add_run((prefix + " " if prefix else "") + text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.1)
    doc.add_paragraph()


def _table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"

    # Header row
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], "1F4E79")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        bg = "F2F7FD" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val) if val is not None else ""
            _set_cell_bg(cells[c_idx], bg)
            for run in cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.color.rgb = DARK_GRAY

    # Column widths
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])

    _set_table_borders(t)
    doc.add_paragraph()
    return t


# ── Document sections ─────────────────────────────────────────────────────────

def _cover(doc, spec, target_system):
    """Full cover page."""
    # Logo (optional — skip if not found)
    logo = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                        "logos", "TechStone.png")
    if os.path.isfile(logo):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo, width=Inches(1.8))
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Migration Design Document")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Subtitle
    source = spec.get("source_system", "Source").upper()
    target = (target_system or "Target").upper()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{source}  →  {target}")
    run2.font.size = Pt(16)
    run2.font.color.rgb = BLUE

    doc.add_paragraph()

    # Metadata table
    project = spec.get("project_name", "Unknown Project")
    summary = spec.get("summary", {})
    flows   = summary.get("total_flows", len(spec.get("integrations", [])))
    gaps    = summary.get("gaps_found", len(spec.get("gaps", [])))
    today   = date.today().strftime("%B %d, %Y")

    meta_rows = [
        ("Project",          project),
        ("Source Platform",  source.title()),
        ("Target Platform",  target.title()),
        ("Total Flows",      str(flows)),
        ("Total Gaps",       str(gaps)),
        ("Document Date",    today),
        ("Prepared By",      "TechStone LLC — MigrAlte"),
    ]
    t = doc.add_table(rows=len(meta_rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(meta_rows):
        cells = t.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        _set_cell_bg(cells[0], "1F4E79")
        _set_cell_bg(cells[1], "F2F7FD" if i % 2 == 0 else "FFFFFF")
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10.5)
        for run in cells[1].paragraphs[0].runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = DARK_GRAY
        cells[0].width = Inches(2.0)
        cells[1].width = Inches(3.5)
    _set_table_borders(t, "D0D8E4")

    doc.add_page_break()


def _section_overview(doc, spec, target_system):
    _h1(doc, "1. Architecture Overview")
    source  = spec.get("source_system", "Unknown")
    target  = target_system or "Unknown"
    summary = spec.get("summary", {})
    notes   = spec.get("migration_notes", "")

    _body(doc, f"This document describes the migration of integrations from "
               f"{source.upper()} to {target.upper()}.")

    if notes:
        _callout(doc, notes, bg_hex="DDEEFF", prefix="Note:")

    # Summary stats table
    stats = [
        ("Total Flows",       summary.get("total_flows",    len(spec.get("integrations", [])))),
        ("Primary Flows",     summary.get("primary_flows",  "—")),
        ("Sub-Flows",         summary.get("sub_flows",       "—")),
        ("Total Connections", summary.get("total_connections", len(spec.get("connections", {})))),
        ("Gaps Found",        summary.get("gaps_found",      len(spec.get("gaps", [])))),
        ("Overall Complexity",summary.get("overall_complexity", "—")),
    ]
    _table(doc, ["Metric", "Value"],
           [[k, str(v)] for k, v in stats],
           col_widths=[2.5, 2.0])


def _section_flows(doc, spec):
    _h1(doc, "2. Integration Flows")
    integrations = spec.get("integrations", [])

    if not integrations:
        _body(doc, "No integration flows found in spec.")
        return

    for idx, flow in enumerate(integrations, 1):
        name       = flow.get("name", f"Flow {idx}")
        flow_type  = flow.get("flow_type", "primary")
        trigger    = flow.get("trigger", {})
        steps      = flow.get("steps", [])
        error_hdl  = flow.get("error_handling", {})
        conns      = flow.get("connections_used", [])

        _h2(doc, f"2.{idx}  {name}")
        _body(doc, f"Type: {flow_type.title()}   |   Steps: {len(steps)}   |   "
                   f"Error Handling: {'Yes' if error_hdl.get('has_error_handler') else 'No'}")

        # Trigger
        _h3(doc, "Trigger")
        t_type  = trigger.get("type", "unknown")
        t_label = trigger.get("label", "")
        t_note  = trigger.get("note", "")
        _bullet(doc, f"Type: {t_type}")
        if t_label:
            _bullet(doc, f"Label: {t_label}")
        if t_note:
            _bullet(doc, f"Note: {t_note}")
        if trigger.get("requires_review"):
            _callout(doc, "Trigger requires manual review before deployment.",
                     bg_hex="FFF2CC", prefix="")

        # Steps table
        if steps:
            _h3(doc, "Steps")
            step_rows = []
            for s in steps:
                seq        = s.get("sequence", "")
                s_type     = s.get("type", "")
                s_label    = s.get("label", "")
                complexity = s.get("complexity", "")
                review     = "Yes" if s.get("requires_review") else ""
                step_rows.append([str(seq), s_type, s_label, complexity, review])

            _table(doc,
                   ["#", "Type", "Label", "Complexity", "Review?"],
                   step_rows,
                   col_widths=[0.4, 1.4, 2.4, 1.0, 0.8])

        # Connections used
        if conns:
            _h3(doc, "Connections Used")
            for c in conns:
                _bullet(doc, c)


def _section_connections(doc, spec):
    _h1(doc, "3. Connections & Dependencies")
    connections = spec.get("connections", {})

    if not connections:
        _body(doc, "No connections defined in spec.")
        return

    rows = []
    for key, conn in connections.items():
        if isinstance(conn, dict):
            conn_type = conn.get("type", conn.get("connector_type", ""))
            host      = conn.get("host", conn.get("url", conn.get("base_url", "")))
            note      = conn.get("note", "")
            rows.append([key, conn_type, host, note])
        else:
            rows.append([key, str(conn), "", ""])

    if rows:
        _table(doc, ["Connection Name", "Type", "Host / URL", "Notes"],
               rows, col_widths=[1.8, 1.4, 2.0, 1.4])
    else:
        _body(doc, "Connection details not available in spec.")


def _section_gaps(doc, spec):
    _h1(doc, "4. Known Gaps & Risks")
    gaps = spec.get("gaps", [])

    if not gaps:
        _callout(doc, "No gaps found — all flows are expected to migrate cleanly.",
                 bg_hex="E2EFDA", prefix="")
        return

    _body(doc, f"{len(gaps)} gap(s) identified during analysis. "
               "Items marked HIGH severity require manual resolution before go-live.")

    rows = []
    for g in gaps:
        flow      = g.get("flow", g.get("integration", ""))
        component = g.get("component", g.get("step_type", ""))
        severity  = g.get("severity", "medium").lower()
        desc      = g.get("description", g.get("message", ""))
        rows.append([flow, component, severity.upper(), desc])

    _table(doc, ["Flow", "Component", "Severity", "Description"],
           rows, col_widths=[1.5, 1.4, 0.9, 2.8])

    # Severity summary
    high_count = sum(1 for g in gaps if g.get("severity", "").lower() == "high")
    if high_count:
        _callout(doc,
                 f"{high_count} HIGH severity gap(s) require manual implementation — "
                 "see rows above.",
                 bg_hex="FFDDC1", prefix="Action Required:")


def _section_next_steps(doc, spec, target_system):
    _h1(doc, "5. Implementation Guide")
    source = spec.get("source_system", "source")
    target = target_system or "target"

    steps = [
        f"Review all HIGH severity gaps in Section 4 and plan manual implementations.",
        f"Configure connections on the {target.upper()} platform "
        f"(see Section 3 for connection details).",
        f"Deploy generated {target.upper()} components to a non-production environment.",
        f"Run functional tests for each flow listed in Section 2.",
        f"Validate error handling and retry behaviour end-to-end.",
        f"Obtain sign-off from business stakeholders before promoting to production.",
        f"Decommission {source.upper()} flows after successful parallel-run validation.",
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  {step}")
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)


def _section_md_appendix(doc, md_dir):
    """Optional appendix: include Markdown analysis files as plain text."""
    if not md_dir or not os.path.isdir(md_dir):
        return

    md_files = sorted(Path(md_dir).glob("*.md"))
    if not md_files:
        return

    doc.add_page_break()
    _h1(doc, "Appendix: Analysis Documents")

    for md_path in md_files:
        _h2(doc, md_path.stem.replace("_", " "))
        try:
            content = md_path.read_text(encoding="utf-8")
            for line in content.splitlines()[:80]:  # cap at 80 lines per file
                stripped = line.strip()
                if stripped.startswith("# "):
                    _h3(doc, stripped[2:])
                elif stripped.startswith("## "):
                    _body(doc, stripped[3:], bold=True)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    _bullet(doc, stripped[2:])
                elif stripped:
                    _body(doc, stripped)
        except Exception:
            _body(doc, f"(Could not read {md_path.name})")


# ── Document setup ────────────────────────────────────────────────────────────

def _setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    return doc


# ── Entry point ───────────────────────────────────────────────────────────────

def generate(spec_path: str, output_path: str, target_system: str = None,
             md_dir: str = None) -> str:
    with open(spec_path, encoding="utf-8") as f:
        spec = json.load(f)

    doc = _setup_doc()

    _cover(doc, spec, target_system)
    _section_overview(doc, spec, target_system)
    _section_flows(doc, spec)
    _section_connections(doc, spec)
    _section_gaps(doc, spec)
    _section_next_steps(doc, spec, target_system)
    _section_md_appendix(doc, md_dir)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Generate a migration design document from a spec JSON.")
    parser.add_argument("spec", help="Path to migration spec JSON file")
    parser.add_argument("--output", "-o", default=None,
                        help="Output .docx path (default: migration-specs/<project>_design_document.docx)")
    parser.add_argument("--target", default=None,
                        help="Target platform name (e.g. workato, boomi) for display in doc")
    parser.add_argument("--md-dir", default=None,
                        help="Optional directory of .md files to include as appendix")
    args = parser.parse_args()

    if not os.path.isfile(args.spec):
        print(f"ERROR: Spec file not found: {args.spec}", file=sys.stderr)
        sys.exit(1)

    # Default output path alongside the spec
    if args.output:
        out_path = args.output
    else:
        spec_stem = Path(args.spec).stem
        out_path  = os.path.join(os.path.dirname(args.spec), f"{spec_stem}_design_document.docx")

    print(f"  Generating Word document: {out_path}")
    result = generate(args.spec, out_path, target_system=args.target, md_dir=args.md_dir)
    print(f"  Done: {result}")


if __name__ == "__main__":
    main()
